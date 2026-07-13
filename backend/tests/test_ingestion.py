"""Tests for ingestion: extraction, chunking provenance, importance, outline."""
from __future__ import annotations

import io

import pytest

from app.ingestion import (
    chunk_blocks,
    chunk_blocks_semantic,
    extract_text,
    sniff_content_ok,
    split_sentences,
)
from app.services import extractor
from app.services.importance import extract_highlights, score_chunks
from app.services.outliner import extract_outline
from ._helpers import make_encrypted_pdf, make_text_pdf


def _make_docx(paras: list[tuple[str, str | None]]) -> bytes:
    """Build a .docx from (text, style) pairs; style None => body paragraph."""
    import docx

    d = docx.Document()
    for text, style in paras:
        d.add_paragraph(text, style=style) if style else d.add_paragraph(text)
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def _make_xlsx(rows: list[list]) -> bytes:
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    for row in rows:
        ws.append(row)
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def test_extract_text_from_pdf():
    pdf = make_text_pdf("Hello PaperTrail ingestion test")
    text, pages = extract_text(pdf, "pdf")
    assert "Hello PaperTrail ingestion test" in text
    assert pages == 1


def test_extract_text_from_txt():
    text, pages = extract_text(b"plain text body", "txt")
    assert text == "plain text body"
    assert pages is None


def test_extract_encrypted_pdf_raises():
    with pytest.raises(Exception):
        extract_text(make_encrypted_pdf(), "pdf")


def test_extract_corrupt_pdf_raises():
    with pytest.raises(Exception):
        extract_text(b"%PDF-1.4 not really a pdf \xde\xad", "pdf")


def test_sniff_valid_pdf():
    assert sniff_content_ok(make_text_pdf("x"), "pdf") is True


def test_sniff_rejects_non_pdf_with_pdf_extension():
    assert sniff_content_ok(b"just text, not a pdf", "pdf") is False


def test_sniff_accepts_utf8_text():
    assert sniff_content_ok("café — résumé".encode("utf-8"), "txt") is True


def test_sniff_rejects_binary_as_text():
    assert sniff_content_ok(b"\x00\x01\x02\x03" * 100, "txt") is False


# --------------------------- extractor blocks --------------------------- #
def test_extract_blocks_pdf_has_page_numbers():
    pdf = make_text_pdf("Quarterly revenue grew across every region this year")
    blocks = extractor.extract_blocks(pdf, "pdf")
    assert blocks and blocks[0]["page"] == 1
    assert "Quarterly revenue" in blocks[0]["text"]


def test_extract_blocks_docx_captures_heading_levels():
    data = _make_docx(
        [
            ("Executive Summary", "Heading 1"),
            ("The company performed well this year in all key markets.", None),
            ("Financial Details", "Heading 2"),
            ("Revenue increased by twenty percent year over year overall.", None),
        ]
    )
    blocks = extractor.extract_blocks(data, "docx")
    headings = [(b["heading"], b["level"]) for b in blocks if b["is_heading"]]
    assert ("Executive Summary", 1) in headings
    assert ("Financial Details", 2) in headings


def test_extract_blocks_xlsx_rows_to_text():
    data = _make_xlsx([["Name", "Amount"], ["Widget", 42], ["Gadget", 99]])
    blocks = extractor.extract_blocks(data, "xlsx")
    texts = " ".join(b["text"] for b in blocks)
    assert "Name: Widget" in texts and "Amount: 42" in texts


def test_extract_blocks_csv_rows_to_text():
    data = b"City,Population\nParis,2100000\nBerlin,3700000\n"
    blocks = extractor.extract_blocks(data, "csv")
    joined = " ".join(b["text"] for b in blocks)
    assert "City: Paris" in joined and "Population: 3700000" in joined


def test_extract_blocks_unsupported_raises_415():
    from fastapi import HTTPException

    with pytest.raises(HTTPException) as ei:
        extractor.extract_blocks(b"data", "exe")
    assert ei.value.status_code == 415


def test_validate_content_rejects_renamed_binary():
    # An .exe (MZ header) renamed to .pdf must fail the magic-byte check.
    assert extractor.validate_content(b"MZ\x90\x00", "pdf") is False
    assert extractor.validate_content(make_text_pdf("x"), "pdf") is True


# ----------------------- chunking provenance ---------------------------- #
def test_chunk_blocks_carries_page_and_section():
    blocks = [
        {"text": "OVERVIEW", "page": 1, "heading": "OVERVIEW", "level": 1, "is_heading": True},
        {"text": "Body paragraph about the overview section here.", "page": 1,
         "heading": None, "level": 0, "is_heading": False},
        {"text": "Second page content that continues the discussion.", "page": 2,
         "heading": None, "level": 0, "is_heading": False},
    ]
    chunks = chunk_blocks(blocks, chunk_size=40, overlap=10)
    assert chunks
    assert all("page_number" in c and "section_heading" in c for c in chunks)
    # The heading section carries into following body chunks.
    assert chunks[0]["section_heading"] == "OVERVIEW"


# ------------------- semantic (sentence-boundary) chunking -------------- #
def test_split_sentences_respects_abbreviations():
    text = "Mr. Smith said hello. He then left the building."
    sentences = split_sentences(text)
    assert sentences == ["Mr. Smith said hello.", "He then left the building."]


def test_split_sentences_handles_multiple_abbreviation_styles():
    text = (
        "Dr. Jones met with Prof. Lee vs. the review board. "
        "The meeting, e.g. a routine check-in, ended quickly."
    )
    sentences = split_sentences(text)
    assert len(sentences) == 2
    assert sentences[0].startswith("Dr. Jones met with Prof. Lee vs. the review board.")
    assert sentences[1].startswith("The meeting, e.g. a routine check-in, ended quickly.")


def test_split_sentences_empty_text():
    assert split_sentences("") == []
    assert split_sentences("   ") == []


def test_chunk_blocks_semantic_same_schema_as_chunk_blocks():
    blocks = [
        {"text": "OVERVIEW", "page": 1, "heading": "OVERVIEW", "level": 1, "is_heading": True},
        {
            "text": "This is the first sentence of the overview. This is the second one.",
            "page": 1,
            "heading": None,
            "level": 0,
            "is_heading": False,
        },
        {
            "text": "This sentence lives on the second page. It also ends cleanly.",
            "page": 2,
            "heading": None,
            "level": 0,
            "is_heading": False,
        },
    ]
    character_chunks = chunk_blocks(blocks, chunk_size=40, overlap=10)
    semantic_chunks = chunk_blocks_semantic(blocks, chunk_size=10, overlap_sentences=1)

    assert character_chunks and semantic_chunks
    expected_keys = {"text", "page_number", "section_heading"}
    for c in character_chunks + semantic_chunks:
        assert set(c) == expected_keys
    assert semantic_chunks[0]["section_heading"] == "OVERVIEW"


def test_chunk_blocks_semantic_never_splits_mid_sentence():
    blocks = [
        {
            "text": (
                "Dr. Evans reviewed the quarterly figures carefully. "
                "Revenue grew by twelve percent this year. "
                "Mr. Patel confirmed the numbers independently. "
                "The board approved the budget for next year."
            ),
            "page": 1,
            "heading": None,
            "level": 0,
            "is_heading": False,
        }
    ]
    chunks = chunk_blocks_semantic(blocks, chunk_size=8, overlap_sentences=1)
    assert len(chunks) > 1
    for c in chunks:
        assert c["text"][-1] in ".?!\"'"


def test_chunk_blocks_semantic_overlaps_by_sentences():
    blocks = [
        {
            "text": " ".join(f"This is sentence number {i}." for i in range(10)),
            "page": 1,
            "heading": None,
            "level": 0,
            "is_heading": False,
        }
    ]
    chunks = chunk_blocks_semantic(blocks, chunk_size=15, overlap_sentences=1)
    assert len(chunks) > 1
    # The last sentence(s) of a chunk reappear at the start of the next chunk.
    first_last_sentence = chunks[0]["text"].rsplit(".", 2)[-2].strip() + "."
    assert first_last_sentence in chunks[1]["text"]


def test_chunk_blocks_semantic_empty_blocks_returns_empty():
    assert chunk_blocks_semantic([]) == []
    assert chunk_blocks_semantic([{"text": "", "page": 1, "is_heading": False}]) == []


# ------------------------- importance scoring --------------------------- #
def test_score_chunks_returns_normalized_scores():
    chunks = [
        "The mitochondria is the powerhouse of the cell in biology.",
        "Quarterly revenue and profit margins expanded significantly this year.",
        "Legal disclaimer and standard boilerplate copyright notice text here.",
        "Machine learning models require large labeled training datasets today.",
    ]
    scores = score_chunks(chunks)
    assert len(scores) == len(chunks)
    assert all(0.0 <= s <= 1.0 for s in scores)


def test_score_chunks_edge_cases():
    assert score_chunks([]) == []
    assert score_chunks(["only one chunk"]) == [1.0]


def test_extract_highlights_shape_and_count():
    chunks = [f"Sentence number {i}. Important finding {i} about topic {i}." for i in range(12)]
    scores = score_chunks(chunks)
    highlights = extract_highlights(chunks, scores, n=8)
    assert len(highlights) == 8
    for h in highlights:
        assert set(h) == {"text", "score", "chunk_index"}
        assert 0.0 <= h["score"] <= 1.0


# ---------------------------- outline ----------------------------------- #
def test_extract_outline_from_docx_headings():
    blocks = [
        {"text": "INTRODUCTION", "heading": "INTRODUCTION", "level": 1, "is_heading": True, "page": 1},
        {"text": "Some introductory content in the document body here.", "heading": None,
         "level": 0, "is_heading": False, "page": 1},
    ]
    chunk_texts = ["INTRODUCTION\n\nSome introductory content in the document body here."]
    outline = extract_outline(blocks, chunk_texts)
    assert len(outline) >= 1
    assert outline[0]["heading"] == "INTRODUCTION"
    assert outline[0]["level"] == 1
    assert outline[0]["chunk_index"] == 0
