"""Router coverage for /api/documents: upload validation, listing, pagination."""
from __future__ import annotations

import time

import pytest

import app.routers.documents as documents_router
from ._helpers import (
    make_encrypted_pdf,
    make_multi_page_pdf,
    make_text_pdf,
    upload_bytes,
    upload_text_doc,
)

from conftest import requires_db

pytestmark = requires_db


def test_upload_txt_creates_chunks(client):
    body = upload_text_doc(client, "notes.txt", "hello world " * 200)
    assert body["filename"] == "notes.txt"
    assert body["chunks_created"] >= 1


def test_upload_pdf_extracts_text(client):
    pdf = make_text_pdf("Invoice total is 1234 dollars for consulting services")
    res = upload_bytes(client, "invoice.pdf", pdf, "application/pdf")
    assert res.status_code == 200, res.text
    assert res.json()["chunks_created"] >= 1


def test_upload_unsupported_type_rejected(client):
    res = upload_bytes(client, "malware.exe", b"MZ\x90\x00binary", "application/octet-stream")
    assert res.status_code == 415
    assert "Unsupported file type" in res.json()["error"]["message"]


def test_upload_empty_file_rejected(client):
    res = upload_bytes(client, "empty.txt", b"", "text/plain")
    assert res.status_code == 400
    assert "empty" in res.json()["error"]["message"].lower()


def test_upload_oversized_file_rejected(client, monkeypatch):
    # Shrink the ceiling so the test payload is comfortably "too big".
    # max_upload_bytes is derived from max_upload_mb, so patching the latter
    # drives the check (0.001 MB ~= 1048 bytes < 2048-byte payload).
    monkeypatch.setattr(documents_router.settings, "max_upload_mb", 0.001, raising=False)
    res = upload_bytes(client, "big.txt", b"x" * 2048, "text/plain")
    assert res.status_code == 413
    assert "maximum upload size" in res.json()["error"]["message"].lower()


def test_upload_txt_with_binary_content_rejected(client):
    # A binary blob renamed .txt must be rejected by content sniffing.
    res = upload_bytes(client, "fake.txt", b"\x00\x01\x02\x03\xff\xfe" * 100, "text/plain")
    assert res.status_code == 422
    assert "content does not match" in res.json()["error"]["message"].lower()


def test_upload_pdf_that_is_not_a_pdf_rejected(client):
    res = upload_bytes(client, "fake.pdf", b"this is plainly not a pdf file", "application/pdf")
    assert res.status_code == 422


def test_upload_corrupt_pdf_returns_422(client):
    # Valid magic bytes but unreadable structure -> graceful 422, not 500.
    res = upload_bytes(client, "broken.pdf", b"%PDF-1.4\n garbage \xde\xad\xbe\xef", "application/pdf")
    assert res.status_code == 422


def test_upload_password_protected_pdf_returns_422(client):
    res = upload_bytes(client, "locked.pdf", make_encrypted_pdf(), "application/pdf")
    assert res.status_code == 422


def test_list_documents_newest_first_with_chunk_counts(client):
    upload_text_doc(client, "first.txt", "alpha content here")
    upload_text_doc(client, "second.txt", "beta content here")
    res = client.get("/api/documents")
    assert res.status_code == 200
    docs = res.json()
    assert docs[0]["filename"] == "second.txt"  # newest first
    assert all(d["chunk_count"] >= 1 for d in docs)


def test_list_documents_pagination(client):
    for i in range(5):
        upload_text_doc(client, f"doc{i}.txt", f"content number {i}")
    page = client.get("/api/documents?limit=2&offset=0")
    assert page.status_code == 200
    assert len(page.json()) == 2

    # Invalid pagination params are rejected by validation.
    assert client.get("/api/documents?limit=0").status_code == 422
    assert client.get("/api/documents?limit=999").status_code == 422


# --------------------- Phase 2: enriched ingestion ---------------------- #
def _docx_bytes(paras):
    import io

    import docx

    d = docx.Document()
    for text, style in paras:
        d.add_paragraph(text, style=style) if style else d.add_paragraph(text)
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def _xlsx_bytes(rows):
    import io

    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    for row in rows:
        ws.append(row)
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def test_upload_returns_highlights_and_importance(client):
    body = upload_text_doc(
        client,
        "report.txt",
        "Quarterly revenue rose sharply. "
        "Machine learning adoption accelerated across teams. "
        "Legal boilerplate and standard disclaimers appear here. " * 20,
    )
    assert 1 <= len(body["highlights"]) <= 8
    for h in body["highlights"]:
        assert 0.0 <= h["score"] <= 1.0 and "text" in h
    assert body["word_count"] > 0


def test_upload_docx_extracts_outline(client):
    data = _docx_bytes(
        [
            ("PROJECT OVERVIEW", "Heading 1"),
            ("This document describes the overall project scope and goals.", None),
            ("Milestones", "Heading 2"),
            ("Key milestones are tracked on a quarterly cadence here.", None),
        ]
    )
    res = upload_bytes(client, "plan.docx", data)
    assert res.status_code == 200, res.text
    outline = res.json()["outline"]
    assert any(o["heading"] == "PROJECT OVERVIEW" for o in outline)


def test_upload_xlsx_creates_queryable_chunks(client):
    data = _xlsx_bytes([["Product", "Price"], ["Alpha", 10], ["Beta", 20]])
    res = upload_bytes(client, "prices.xlsx", data)
    assert res.status_code == 200, res.text
    assert res.json()["chunks_created"] >= 1
    assert res.json()["file_type"] == "xlsx"


# --------------- V3 Phase 1: sentence-boundary (semantic) chunking ------- #
def test_upload_pdf_semantic_chunking_no_chunk_ends_mid_sentence(client, db_session):
    from sqlalchemy import select

    from app.models import Chunk, Document

    pages = [
        (
            f"This is page {i} of the annual report. "
            f"Dr. Evans reviewed section {i} carefully and confirmed the figures. "
            f"Revenue in region {i} grew by a healthy margin this year. "
            f"Mr. Patel signed off on the results for this section."
        )
        for i in range(1, 11)
    ]
    pdf = make_multi_page_pdf(pages)
    res = upload_bytes(client, "annual-report.pdf", pdf, "application/pdf")
    assert res.status_code == 200, res.text
    assert res.json()["page_count"] == 10

    doc_id = res.json()["id"]
    doc = db_session.execute(select(Document).where(Document.id == doc_id)).scalar_one()
    assert doc.chunking_strategy == "semantic"

    chunks = db_session.execute(
        select(Chunk.content).where(Chunk.document_id == doc_id)
    ).scalars().all()
    assert chunks
    for content in chunks:
        assert content[-1] in ".?!\"'", f"chunk ends mid-sentence: {content!r}"
        # "Mr." / "Dr." must never be split away from the name that follows.
        assert not content.rstrip().endswith(("Mr.", "Dr."))


def test_upload_txt_character_strategy_recorded(client, db_session, monkeypatch):
    import app.routers.documents as documents_router
    from sqlalchemy import select

    from app.models import Document

    monkeypatch.setattr(documents_router.settings, "chunking_strategy", "character", raising=False)
    body = upload_text_doc(client, "legacy.txt", "hello world " * 200)

    doc = db_session.execute(
        select(Document).where(Document.id == body["id"])
    ).scalar_one()
    assert doc.chunking_strategy == "character"


def test_document_status_reports_processed(client):
    body = upload_text_doc(client, "s.txt", "some content to process here " * 10)
    res = client.get(f"/api/documents/{body['id']}/status")
    assert res.status_code == 200
    payload = res.json()
    assert payload["processed"] is True
    assert payload["chunk_count"] >= 1
    assert payload["processed_at"] is not None


def test_document_status_cross_user_forbidden(client, anon_client):
    from conftest import auth_headers

    body = upload_text_doc(client, "mine.txt", "private content here " * 10)
    other = auth_headers(anon_client, "intruder@papertrail.io")
    res = anon_client.get(f"/api/documents/{body['id']}/status", headers=other)
    assert res.status_code == 403


# ----------------- off-loop processing (event-loop safety) --------------- #
def test_document_status_reports_processing_status_done(client):
    body = upload_text_doc(client, "ok.txt", "perfectly fine content " * 20)
    status = client.get(f"/api/documents/{body['id']}/status").json()
    assert status["processing_status"] == "done"
    assert status["error"] is None


def test_extract_failure_surfaces_failed_status(client, db_session, monkeypatch):
    from sqlalchemy import select

    from app.models import Document

    def boom(data, file_type):
        raise RuntimeError("simulated extractor crash")

    monkeypatch.setattr(documents_router.extractor, "extract_blocks", boom)
    res = upload_bytes(client, "bad.txt", b"plain text content", "text/plain")
    assert res.status_code == 422

    # The row created before processing records the failure for the frontend.
    doc_id = db_session.execute(
        select(Document.id).where(Document.filename == "bad.txt")
    ).scalar_one()
    status = client.get(f"/api/documents/{doc_id}/status").json()
    assert status["processing_status"] == "failed"
    assert status["error"]
    assert status["processed"] is False

    # A failed ingestion never becomes part of the library listing.
    assert doc_id not in {d["id"] for d in client.get("/api/documents").json()}


@pytest.fixture
def anyio_backend():
    return "asyncio"


@pytest.mark.anyio
async def test_concurrent_request_not_blocked_by_upload(db_session, monkeypatch):
    """Prove the heavy pipeline runs off the event loop: /api/health completes
    while an upload is stuck inside a slow, fully blocking embed call."""
    import asyncio

    import httpx

    from app import llm
    from app.database import get_db
    from app.main import app

    def slow_embed(texts):
        time.sleep(0.5)  # blocking, like the real network/CPU-bound call
        return [[0.1, 0.2, 0.3] for _ in texts]

    monkeypatch.setattr(llm, "embed_texts", slow_embed)

    def _override_get_db():
        yield db_session

    app.dependency_overrides[get_db] = _override_get_db
    try:
        transport = httpx.ASGITransport(app=app)
        async with httpx.AsyncClient(transport=transport, base_url="http://test") as ac:
            reg = await ac.post(
                "/api/auth/register",
                json={"email": "loop@papertrail.io", "password": "password123"},
            )
            headers = {"Authorization": f"Bearer {reg.json()['access_token']}"}

            upload = asyncio.create_task(
                ac.post(
                    "/api/documents/upload",
                    files={"file": ("t.txt", b"hello world " * 100, "text/plain")},
                    headers=headers,
                )
            )
            await asyncio.sleep(0.1)  # let the upload reach the blocking embed

            start = time.perf_counter()
            health = await ac.get("/api/health")
            elapsed = time.perf_counter() - start

            res = await upload
    finally:
        app.dependency_overrides.pop(get_db, None)

    assert res.status_code == 200, res.text
    assert health.status_code == 200
    # On the event loop this would wait out the 0.5s embed sleep; off-loop it
    # completes effectively instantly.
    assert elapsed < 0.2, f"/api/health was blocked for {elapsed:.3f}s"
